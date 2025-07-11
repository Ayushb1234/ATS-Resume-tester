import pdfplumber
from docx import Document

def extract_text_from_pdf(path):
    """
    Extracts text from PDF file using pdfplumber.
    """
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text.strip()

def extract_text_from_docx(file):
    """
    Extracts text from a DOCX file (uploaded stream or path).
    """
    try:
        if hasattr(file, "read"):  # If it's an uploaded file (Streamlit)
            doc = Document(file)
        else:
            doc = Document(file)  # Local file path
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""
